from __future__ import annotations

import io
import json
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote
from uuid import uuid4
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import APIRouter, Body, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field

from app.db.session import SessionLocal
from app.services.file_service import FileService

router = APIRouter()


class Step1ExportRequest(BaseModel):
    user_id: str = Field(default="demo-user-id")
    project_name: str = Field(default="")
    thread_id: str = Field(default="")
    content_text: str = Field(default="", min_length=1)
    content_json: str | None = None
    export_style: str = Field(default="classic")
    custom_title: str | None = None
    save_to_database: bool = Field(default=True)
    draft_payload: dict[str, Any] | None = None


class Step2FormatOptions(BaseModel):
    """Typography options for Step 2 custom-style export."""

    font_family: str | None = Field(default=None, description="正文字体（中文字体名）")
    font_size_pt: float | None = Field(default=None, description="正文字号（pt）")
    heading_font_size_pt: float | None = Field(default=None, description="一级标题字号（pt）")
    line_spacing: float | None = Field(default=None, description="行距（倍数，1.0 ~ 3.0）")
    paragraph_spacing_pt: float | None = Field(default=None, description="段后距（pt）")
    first_line_indent_chars: float | None = Field(default=None, description="首行缩进（中文字符数）")


class Step2ExportRequest(BaseModel):
    user_id: str = Field(default="demo-user-id")
    project_name: str = Field(default="")
    thread_id: str = Field(default="")
    content_text: str = Field(default="", min_length=1)
    content_json: str | None = None
    export_style: str = Field(default="classic")
    custom_title: str | None = None
    save_to_database: bool = Field(default=True)
    draft_payload: dict[str, Any] | None = None
    categories: list[str] | None = Field(default=None, description="纳入封面页的分类清单")
    format_options: Step2FormatOptions | None = None


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', value).strip(' ._')
    return cleaned or '未命名项目'


def _attachment_disposition(filename: str) -> str:
    """Build a latin-1-safe Content-Disposition header with UTF-8 filename fallback."""
    ascii_fallback = re.sub(r'[^\x20-\x7E]', '_', filename).strip(' ._') or 'export.docx'
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"


def _paragraph(text: str, style: str = 'Normal', align: str | None = None) -> str:
    style_xml = f'<w:pStyle w:val="{style}"/>' if style else ''
    align_xml = f'<w:jc w:val="{align}"/>' if align else ''
    return f'<w:p><w:pPr>{style_xml}{align_xml}</w:pPr><w:r><w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>'


def _table(rows: list[tuple[str, str]]) -> str:
    row_xml = []
    for key, value in rows:
        row_xml.append(
            '<w:tr>'
            f'<w:tc><w:tcPr><w:tcW w:w="2400" w:type="dxa"/><w:shd w:fill="F2F6FC"/></w:tcPr><w:p><w:r><w:t>{escape(key)}</w:t></w:r></w:p></w:tc>'
            f'<w:tc><w:tcPr><w:tcW w:w="6800" w:type="dxa"/></w:tcPr><w:p><w:r><w:t>{escape(value)}</w:t></w:r></w:p></w:tc>'
            '</w:tr>'
        )
    return '<w:tbl><w:tblPr><w:tblW w:w="9200" w:type="dxa"/><w:tblBorders><w:top w:val="single" w:sz="4"/><w:left w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/><w:insideH w:val="single" w:sz="4"/><w:insideV w:val="single" w:sz="4"/></w:tblBorders></w:tblPr>' + ''.join(row_xml) + '</w:tbl>'


def _render_content_paragraphs(content: str, export_style: str) -> str:
    blocks: list[str] = []
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            blocks.append('<w:p/>')
            continue
        if line.startswith(('# ', '一、', '二、', '三、', '四、', '五、')):
            blocks.append(_paragraph(line.lstrip('#').strip(), 'Heading1'))
        elif line.startswith(('（一）', '（二）', '（三）', '（四）')) or re.match(r'^\d+[.、]', line):
            blocks.append(_paragraph(line, 'Heading2'))
        elif export_style == 'custom' and line.startswith(('-', '•', '*')):
            blocks.append(_paragraph(f'• {line.lstrip("-•* ")}', 'ListParagraph'))
        else:
            blocks.append(_paragraph(line, 'Normal'))
    return ''.join(blocks)


def _document_xml(title: str, content: str, export_style: str, project_name: str) -> str:
    now = datetime.now().strftime('%Y年%m月%d日')
    intro = _table([
        ('项目名称', project_name),
        ('文档类型', '项目资料清单'),
        ('排版样式', '经典排版' if export_style == 'classic' else '自定义排版'),
        ('导出时间', now),
    ])
    cover = ''.join([
        _paragraph(title, 'Title', 'center'),
        _paragraph('项目资料清单成果文件', 'Subtitle', 'center'),
        _paragraph(now, 'Normal', 'center'),
        '<w:p/>',
        intro,
        '<w:p/>',
        _paragraph('一、资料清单正文', 'Heading1'),
    ])
    content_xml = _render_content_paragraphs(content, export_style)
    margins = '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>' if export_style == 'classic' else '<w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1134"/>'
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {cover}{content_xml}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/>{margins}<w:docGrid w:linePitch="312"/></w:sectPr>
  </w:body>
</w:document>'''


def _styles_xml() -> str:
    return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:line="360" w:lineRule="auto"/><w:ind w:firstLine="480"/></w:pPr><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:qFormat/><w:pPr><w:spacing w:before="240" w:after="240"/></w:pPr><w:rPr><w:b/><w:sz w:val="40"/><w:szCs w:val="40"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:qFormat/><w:rPr><w:color w:val="666666"/><w:sz w:val="26"/><w:szCs w:val="26"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:before="240" w:after="120"/></w:pPr><w:rPr><w:b/><w:sz w:val="30"/><w:szCs w:val="30"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:before="160" w:after="80"/></w:pPr><w:rPr><w:b/><w:sz w:val="26"/><w:szCs w:val="26"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="480" w:firstLine="0"/></w:pPr><w:rPr><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:style>
</w:styles>'''


def _write_docx(path: Path, title: str, content: str, export_style: str, project_name: str) -> None:
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as docx:
        docx.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>''')
        docx.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>''')
        docx.writestr('word/_rels/document.xml.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>''')
        docx.writestr('word/styles.xml', _styles_xml())
        docx.writestr('word/document.xml', _document_xml(title, content, export_style, project_name))


def _clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def _step2_typography(style: str, options: Step2FormatOptions | None) -> dict[str, Any]:
    """Resolve font / size / spacing values, mixing defaults and overrides."""
    base = {
        'font_family': 'SimSun',
        'font_size_pt': 12.0,
        'heading_font_size_pt': 16.0,
        'line_spacing': 1.5,
        'paragraph_spacing_pt': 6.0,
        'first_line_indent_chars': 2.0,
    }
    if style == 'custom' and options is not None:
        data = options.model_dump(exclude_none=True)
        for key, value in data.items():
            if isinstance(value, str) and value.strip():
                base[key] = value.strip()
            elif isinstance(value, (int, float)):
                base[key] = float(value)
    base['font_size_pt'] = _clamp(float(base['font_size_pt']), 8.0, 36.0)
    base['heading_font_size_pt'] = _clamp(float(base['heading_font_size_pt']), 10.0, 48.0)
    base['line_spacing'] = _clamp(float(base['line_spacing']), 1.0, 3.0)
    base['paragraph_spacing_pt'] = _clamp(float(base['paragraph_spacing_pt']), 0.0, 36.0)
    base['first_line_indent_chars'] = _clamp(float(base['first_line_indent_chars']), 0.0, 8.0)
    return base


def _half_points(pt: float) -> int:
    return max(1, int(round(pt * 2)))


def _line_spacing_twips(multiplier: float) -> int:
    return max(240, int(round(multiplier * 240)))


def _twips_from_pt(pt: float) -> int:
    return max(0, int(round(pt * 20)))


def _indent_twips_from_chars(chars: float, font_size_pt: float) -> int:
    # 1 Chinese char ≈ 1 font size em ≈ font_size_pt points, 20 twips per point
    return max(0, int(round(chars * font_size_pt * 20)))


def _step2_paragraph(text: str, *, style: str = '', align: str | None = None, indent_twips: int | None = None, run_attrs: str = '') -> str:
    pieces: list[str] = []
    if style:
        pieces.append(f'<w:pStyle w:val="{style}"/>')
    if align:
        pieces.append(f'<w:jc w:val="{align}"/>')
    if indent_twips is not None:
        pieces.append(f'<w:ind w:firstLine="{indent_twips}"/>')
    p_pr = ''.join(pieces)
    body = f'<w:r>{run_attrs}<w:t xml:space="preserve">{escape(text)}</w:t></w:r>'
    return f'<w:p><w:pPr>{p_pr}</w:pPr>{body}</w:p>'


def _render_step2_paragraphs(content: str, typography: dict[str, Any]) -> str:
    blocks: list[str] = []
    body_indent = _indent_twips_from_chars(float(typography['first_line_indent_chars']), float(typography['font_size_pt']))
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            blocks.append('<w:p/>')
            continue
        stripped = line.strip()
        # Markdown headings
        if stripped.startswith('### '):
            blocks.append(_step2_paragraph(stripped[4:].strip(), style='Heading2'))
            continue
        if stripped.startswith('## '):
            blocks.append(_step2_paragraph(stripped[3:].strip(), style='Heading1'))
            continue
        if stripped.startswith('# '):
            blocks.append(_step2_paragraph(stripped[2:].strip(), style='Title', align='center'))
            continue
        # Numbered or Chinese-numbered headings
        if re.match(r'^(一|二|三|四|五|六|七|八|九|十)[、.]', stripped):
            blocks.append(_step2_paragraph(stripped, style='Heading1'))
            continue
        if re.match(r'^（[一二三四五六七八九十]+）', stripped) or re.match(r'^\d+[.、]', stripped):
            blocks.append(_step2_paragraph(stripped, style='Heading2'))
            continue
        if stripped.startswith(('- ', '• ', '* ')):
            blocks.append(_step2_paragraph(f'• {stripped.lstrip("-•* ")}', style='ListParagraph', indent_twips=body_indent))
            continue
        blocks.append(_step2_paragraph(stripped, style='Normal', indent_twips=body_indent))
    return ''.join(blocks)


def _step2_document_xml(
    *,
    title: str,
    content: str,
    export_style: str,
    project_name: str,
    categories: list[str] | None,
    typography: dict[str, Any],
) -> str:
    now = datetime.now().strftime('%Y年%m月%d日')
    category_text = '、'.join([c for c in (categories or []) if c]) or '资金管理类、预算管理类、制度文件类、项目实施类'
    intro = _table([
        ('项目名称', project_name),
        ('文档类型', '项目核心内容'),
        ('排版样式', '经典排版' if export_style == 'classic' else '自定义排版'),
        ('分类清单', category_text),
        ('字体设置', f"{typography['font_family']} / 正文 {typography['font_size_pt']:.1f}pt / 标题 {typography['heading_font_size_pt']:.1f}pt"),
        ('行距 / 段后距', f"{typography['line_spacing']:.2f} 倍 / {typography['paragraph_spacing_pt']:.1f}pt"),
        ('导出时间', now),
    ])
    cover = ''.join([
        _step2_paragraph(title, style='Title', align='center'),
        _step2_paragraph('项目核心内容成果文件', style='Subtitle', align='center'),
        _step2_paragraph(now, style='Normal', align='center'),
        '<w:p/>',
        intro,
        '<w:p/>',
        _step2_paragraph('一、核心内容正文', style='Heading1'),
    ])
    body_xml = _render_step2_paragraphs(content, typography)
    margins = (
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>'
        if export_style == 'classic'
        else '<w:pgMar w:top="1134" w:right="1134" w:bottom="1134" w:left="1134"/>'
    )
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {cover}{body_xml}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/>{margins}<w:docGrid w:linePitch="312"/></w:sectPr>
  </w:body>
</w:document>'''


def _step2_styles_xml(typography: dict[str, Any]) -> str:
    font = escape(str(typography['font_family']))
    body_size = _half_points(float(typography['font_size_pt']))
    heading_size = _half_points(float(typography['heading_font_size_pt']))
    title_size = _half_points(float(typography['heading_font_size_pt']) + 6.0)
    subtitle_size = _half_points(float(typography['font_size_pt']) + 1.0)
    list_size = body_size
    line_twips = _line_spacing_twips(float(typography['line_spacing']))
    para_after_twips = _twips_from_pt(float(typography['paragraph_spacing_pt']))
    indent_twips = _indent_twips_from_chars(float(typography['first_line_indent_chars']), float(typography['font_size_pt']))
    font_attr = f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}" w:cs="{font}"/>'
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:docDefaults><w:rPrDefault><w:rPr>{font_attr}<w:sz w:val="{body_size}"/><w:szCs w:val="{body_size}"/></w:rPr></w:rPrDefault></w:docDefaults>
  <w:style w:type="paragraph" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:line="{line_twips}" w:lineRule="auto" w:after="{para_after_twips}"/><w:ind w:firstLine="{indent_twips}"/></w:pPr><w:rPr>{font_attr}<w:sz w:val="{body_size}"/><w:szCs w:val="{body_size}"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:qFormat/><w:pPr><w:spacing w:before="240" w:after="240"/></w:pPr><w:rPr>{font_attr}<w:b/><w:sz w:val="{title_size}"/><w:szCs w:val="{title_size}"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:qFormat/><w:rPr>{font_attr}<w:color w:val="666666"/><w:sz w:val="{subtitle_size}"/><w:szCs w:val="{subtitle_size}"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:before="240" w:after="120"/><w:ind w:firstLine="0"/></w:pPr><w:rPr>{font_attr}<w:b/><w:sz w:val="{heading_size}"/><w:szCs w:val="{heading_size}"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:next w:val="Normal"/><w:qFormat/><w:pPr><w:spacing w:before="160" w:after="80"/><w:ind w:firstLine="0"/></w:pPr><w:rPr>{font_attr}<w:b/><w:sz w:val="{max(heading_size - 4, body_size + 2)}"/><w:szCs w:val="{max(heading_size - 4, body_size + 2)}"/></w:rPr></w:style>
  <w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="480" w:firstLine="0"/></w:pPr><w:rPr>{font_attr}<w:sz w:val="{list_size}"/><w:szCs w:val="{list_size}"/></w:rPr></w:style>
</w:styles>'''


def _write_step2_docx(
    path: Path,
    *,
    title: str,
    content: str,
    export_style: str,
    project_name: str,
    categories: list[str] | None,
    typography: dict[str, Any],
) -> None:
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as docx:
        docx.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>''')
        docx.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>''')
        docx.writestr('word/_rels/document.xml.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>''')
        docx.writestr('word/styles.xml', _step2_styles_xml(typography))
        docx.writestr('word/document.xml', _step2_document_xml(
            title=title,
            content=content,
            export_style=export_style,
            project_name=project_name,
            categories=categories,
            typography=typography,
        ))


@router.post('/step1/{project_id}')
def export_step1(project_id: str, payload: Step1ExportRequest = Body(...)) -> dict[str, Any]:
    content = payload.content_text.strip()
    if not content:
        raise HTTPException(status_code=400, detail='content_text is required')

    project_name = payload.project_name.strip() or f'项目 {project_id}'
    style = payload.export_style if payload.export_style in {'classic', 'custom'} else 'classic'
    style_label = '经典排版' if style == 'classic' else '自定义排版'
    title = payload.custom_title.strip() if payload.custom_title and payload.custom_title.strip() else f'{project_name}项目资料清单'
    filename = f'{_safe_filename(project_name)}项目资料清单_{style_label}.docx'
    export_dir = Path(__file__).resolve().parents[4] / 'storage' / 'projects' / project_id / 'exports'
    export_dir.mkdir(parents=True, exist_ok=True)
    export_path = export_dir / f'{uuid4().hex}_{filename}'
    placeholder_values = {'最终版本内容将在这里编辑。', '最后版本内容将在这里编辑。'}
    if content in placeholder_values or len(content) < 10:
        raise HTTPException(status_code=400, detail='当前成品内容为空或仍为占位文本，请先生成/编辑 Step1 成果后再导出')
    _write_docx(export_path, title, content, style, project_name)

    metadata = {
        'project_id': project_id,
        'project_name': project_name,
        'thread_id': payload.thread_id,
        'content_text': content,
        'content_json': payload.content_json,
        'export_style': style,
        'export_filename': filename,
        'storage_key': str(export_path),
        'exported_at': datetime.utcnow().isoformat(),
        'draft_payload': payload.draft_payload,
    }
    file_record: dict[str, Any] | None = None
    if payload.save_to_database:
        with SessionLocal() as db:
            service = FileService(db)
            file_record = service.create_file_record(
                project_id=project_id,
                user_id=payload.user_id,
                project_name=project_name,
                file_name=filename,
                file_type='docx',
                storage_key=str(export_path),
                source_type='step1_export_final',
                file_size=export_path.stat().st_size,
                metadata_json=json.dumps(metadata, ensure_ascii=False),
                draft_thread_id=payload.thread_id or None,
                draft_payload=payload.draft_payload,
            )

    return {
        'project_id': project_id,
        'file_name': filename,
        'storage_key': str(export_path),
        'download_url': f'/api/v1/exports/download/{project_id}/{quote(export_path.name, safe="")}',
        'file_record': file_record,
        'metadata': metadata,
    }


@router.post('/step2/{project_id}')
def export_step2(project_id: str, payload: Step2ExportRequest = Body(...)) -> dict[str, Any]:
    content = payload.content_text.strip()
    if not content:
        raise HTTPException(status_code=400, detail='content_text is required')
    placeholder_values = {'最终版本内容将在这里编辑。', '最后版本内容将在这里编辑。'}
    if content in placeholder_values or len(content) < 10:
        raise HTTPException(status_code=400, detail='当前核心内容为空或仍为占位文本，请先生成 / 编辑 Step2 成果后再导出')

    project_name = payload.project_name.strip() or f'项目 {project_id}'
    style = payload.export_style if payload.export_style in {'classic', 'custom'} else 'classic'
    style_label = '经典排版' if style == 'classic' else '自定义排版'
    title = payload.custom_title.strip() if payload.custom_title and payload.custom_title.strip() else f'{project_name}项目核心内容'
    filename = f'{_safe_filename(project_name)}项目核心内容_{style_label}.docx'
    export_dir = Path(__file__).resolve().parents[4] / 'storage' / 'projects' / project_id / 'exports'
    export_dir.mkdir(parents=True, exist_ok=True)
    export_path = export_dir / f'{uuid4().hex}_{filename}'

    typography = _step2_typography(style, payload.format_options)
    categories = [c for c in (payload.categories or []) if isinstance(c, str) and c.strip()]

    _write_step2_docx(
        export_path,
        title=title,
        content=content,
        export_style=style,
        project_name=project_name,
        categories=categories,
        typography=typography,
    )

    metadata = {
        'project_id': project_id,
        'project_name': project_name,
        'thread_id': payload.thread_id,
        'content_text': content,
        'content_json': payload.content_json,
        'export_style': style,
        'export_filename': filename,
        'storage_key': str(export_path),
        'exported_at': datetime.utcnow().isoformat(),
        'draft_payload': payload.draft_payload,
        'categories': categories,
        'typography': typography,
    }
    file_record: dict[str, Any] | None = None
    if payload.save_to_database:
        with SessionLocal() as db:
            service = FileService(db)
            file_record = service.create_file_record(
                project_id=project_id,
                user_id=payload.user_id,
                project_name=project_name,
                file_name=filename,
                file_type='docx',
                storage_key=str(export_path),
                source_type='step2_export_final',
                file_size=export_path.stat().st_size,
                metadata_json=json.dumps(metadata, ensure_ascii=False),
                draft_thread_id=payload.thread_id or None,
                draft_payload=payload.draft_payload,
            )

    return {
        'project_id': project_id,
        'file_name': filename,
        'storage_key': str(export_path),
        'download_url': f'/api/v1/exports/download/{project_id}/{quote(export_path.name, safe="")}',
        'file_record': file_record,
        'metadata': metadata,
    }


@router.get('/download/{project_id}/{filename}')
def download_export(project_id: str, filename: str) -> FileResponse:
    safe_name = Path(filename).name
    export_path = Path(__file__).resolve().parents[4] / 'storage' / 'projects' / project_id / 'exports' / safe_name
    if not export_path.exists():
        raise HTTPException(status_code=404, detail='export file not found')
    display_name = safe_name.split('_', 1)[-1] if '_' in safe_name else safe_name
    return FileResponse(
        path=export_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': _attachment_disposition(display_name)},
    )


STEP_TITLE_MAP: dict[str, str] = {
    'step1': '项目资料清单',
    'step2': '有效项目资料',
    'step3': '指标体系',
    'step4': '生成分值',
    'step5': '评分标准',
    'step14': '评价报告',
}


class GenericStepExportRequest(BaseModel):
    project_name: str = Field(default='')
    content_text: str = Field(default='', min_length=1)
    export_style: str = Field(default='classic')
    custom_title: str | None = None
    user_id: str = Field(default='demo-user-id')
    project_id: str | None = None
    save_to_database: bool = Field(default=True)


def _step_default_title(step_code: str) -> str:
    if step_code in STEP_TITLE_MAP:
        return f'{step_code.capitalize()} · {STEP_TITLE_MAP[step_code]}'
    digits = ''.join(ch for ch in step_code if ch.isdigit())
    return f'Step{digits or step_code} · 阶段成果' if digits else f'{step_code} · 阶段成果'


def _build_generic_docx(
    *,
    title: str,
    content: str,
    export_style: str,
    project_name: str,
    step_code: str,
) -> bytes:
    doc = Document()

    is_classic = export_style == 'classic'
    body_size = 12 if is_classic else 11
    heading_size = 18 if is_classic else 16

    normal = doc.styles['Normal']
    normal.font.name = 'SimSun' if is_classic else 'Microsoft YaHei'
    normal.font.size = Pt(body_size)

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = cover_title.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(heading_size + 6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f'{STEP_TITLE_MAP.get(step_code, "阶段成果")}成果文件')
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(body_size + 2)

    now_text = datetime.now().strftime('%Y年%m月%d日')
    timestamp_para = doc.add_paragraph()
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_para.add_run(now_text)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1' if 'Light Grid Accent 1' in [s.name for s in doc.styles] else 'Table Grid'
    rows_data = [
        ('项目名称', project_name),
        ('阶段编号', step_code),
        ('排版样式', '经典排版' if is_classic else '自定义排版'),
        ('导出时间', now_text),
    ]
    for row_idx, (key, value) in enumerate(rows_data):
        info_table.rows[row_idx].cells[0].text = key
        info_table.rows[row_idx].cells[1].text = value

    doc.add_paragraph()
    section_heading = doc.add_heading(f'一、{STEP_TITLE_MAP.get(step_code, "阶段成果")}正文', level=1)
    for run in section_heading.runs:
        run.font.size = Pt(heading_size)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith('### '):
            heading = doc.add_heading(stripped[4:].strip(), level=3)
            for run in heading.runs:
                run.font.size = Pt(max(body_size + 2, heading_size - 4))
            continue
        if stripped.startswith('## '):
            heading = doc.add_heading(stripped[3:].strip(), level=2)
            for run in heading.runs:
                run.font.size = Pt(max(body_size + 4, heading_size - 2))
            continue
        if stripped.startswith('# '):
            heading = doc.add_heading(stripped[2:].strip(), level=1)
            for run in heading.runs:
                run.font.size = Pt(heading_size)
            continue
        if re.match(r'^(一|二|三|四|五|六|七|八|九|十)[、.]', stripped):
            heading = doc.add_heading(stripped, level=1)
            for run in heading.runs:
                run.font.size = Pt(heading_size)
            continue
        if re.match(r'^（[一二三四五六七八九十]+）', stripped) or re.match(r'^\d+[.、]', stripped):
            heading = doc.add_heading(stripped, level=2)
            for run in heading.runs:
                run.font.size = Pt(max(body_size + 2, heading_size - 4))
            continue
        if stripped.startswith(('- ', '• ', '* ')):
            bullet = doc.add_paragraph(stripped.lstrip('-•* '), style='List Bullet')
            for run in bullet.runs:
                run.font.size = Pt(body_size)
            continue
        para = doc.add_paragraph(stripped)
        for run in para.runs:
            run.font.size = Pt(body_size)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.post('/{step_code}/word')
def export_step_word(step_code: str, payload: GenericStepExportRequest = Body(...)) -> StreamingResponse:
    code = step_code.lower().strip()
    if not re.match(r'^step([3-9]|1[0-4])$', code):
        raise HTTPException(status_code=400, detail='step_code 必须是 step3~step14（step1/step2 请使用专用导出接口）')

    content = payload.content_text.strip()
    if not content:
        raise HTTPException(status_code=400, detail='content_text is required')
    placeholder_values = {'最终版本内容将在这里编辑。', '最后版本内容将在这里编辑。'}
    if content in placeholder_values or len(content) < 10:
        raise HTTPException(status_code=400, detail='当前成品内容为空或仍为占位文本，请先生成 / 编辑后再导出')

    project_name = payload.project_name.strip() or '未命名项目'
    style = payload.export_style if payload.export_style in {'classic', 'custom'} else 'classic'
    title = (
        payload.custom_title.strip()
        if payload.custom_title and payload.custom_title.strip()
        else _step_default_title(code)
    )

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'{_safe_filename(project_name)}_{code}_{timestamp}.docx'

    docx_bytes = _build_generic_docx(
        title=title,
        content=content,
        export_style=style,
        project_name=project_name,
        step_code=code,
    )

    project_id = (payload.project_id or '').strip()
    if payload.save_to_database and project_id:
        export_dir = Path(__file__).resolve().parents[4] / 'storage' / 'projects' / project_id / 'exports'
        export_dir.mkdir(parents=True, exist_ok=True)
        export_path = export_dir / f'{uuid4().hex}_{filename}'
        export_path.write_bytes(docx_bytes)

        metadata = {
            'project_id': project_id,
            'project_name': project_name,
            'step_code': code,
            'content_text': content,
            'export_style': style,
            'export_filename': filename,
            'storage_key': str(export_path),
            'exported_at': datetime.utcnow().isoformat(),
        }
        with SessionLocal() as db:
            FileService(db).create_file_record(
                project_id=project_id,
                user_id=payload.user_id,
                project_name=project_name,
                file_name=filename,
                file_type='docx',
                storage_key=str(export_path),
                source_type=f'{code}_export_final',
                file_size=len(docx_bytes),
                metadata_json=json.dumps(metadata, ensure_ascii=False),
            )

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': _attachment_disposition(filename)},
    )
